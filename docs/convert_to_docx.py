"""
Convert FSD - CASHPOS PHASE 2.md to .docx using python-docx.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_FILE = "FSD - CASHPOS PHASE 2.md"
DOCX_FILE = "FSD - CASHPOS PHASE 2.docx"


def set_cell_background(cell, fill_color: str):
    """Set table cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table starting at start_idx, add to doc, return end_idx."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|"):
            table_lines.append(line)
            i += 1
        else:
            break

    if not table_lines:
        return start_idx

    # Filter out separator rows (|---|---|)
    data_rows = [r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r)]
    if not data_rows:
        return i

    # Parse cells
    def parse_row(row):
        cells = [c.strip() for c in row.strip("|").split("|")]
        return cells

    rows = [parse_row(r) for r in data_rows]
    col_count = max(len(r) for r in rows)

    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        tr = table.add_row()
        for col_idx in range(col_count):
            cell = tr.cells[col_idx]
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            # Strip inline markdown bold
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"`(.*?)`", r"\1", text)
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_background(cell, "BDD7EE")  # light blue header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

    return i


def apply_inline_formatting(para, text):
    """Add runs with bold/italic/code inline formatting."""
    # Regex to split on **bold**, *italic*, `code`
    pattern = r"(\*\*.*?\*\*|`.*?`|\*.*?\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            if part:
                para.add_run(part)


def main():
    with open(MD_FILE, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading colors
    heading_colors = {
        1: RGBColor(0x1F, 0x49, 0x7D),  # dark blue
        2: RGBColor(0x2E, 0x74, 0xB5),  # medium blue
        3: RGBColor(0x5B, 0x9B, 0xD5),  # lighter blue
        4: RGBColor(0x2F, 0x52, 0x96),  # dark medium
        5: RGBColor(0x40, 0x40, 0x40),  # dark gray
        6: RGBColor(0x60, 0x60, 0x60),
    }

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Code block toggle
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # End code block — add as styled paragraph
                in_code_block = False
                if code_lines:
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.left_indent = Inches(0.25)
                    # Gray shading for paragraph
                    pPr = para._p.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F2F2F2")
                    pPr.append(shd)
                    run = para.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(8.5)
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", stripped):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading_level = min(level, 6)
            para = doc.add_heading(text, level=heading_level)
            for run in para.runs:
                run.font.color.rgb = heading_colors.get(heading_level, RGBColor(0, 0, 0))
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            i = add_table_from_md(doc, lines, i)
            continue

        # Unordered list
        m = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            text = m.group(2)
            level = indent // 2
            style_name = "List Bullet" if level == 0 else "List Bullet 2"
            para = doc.add_paragraph(style=style_name)
            apply_inline_formatting(para, text)
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            text = m.group(2)
            para = doc.add_paragraph(style="List Number")
            apply_inline_formatting(para, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        apply_inline_formatting(para, stripped)
        i += 1

    doc.save(DOCX_FILE)
    print(f"Saved: {DOCX_FILE}")

    # Stats
    doc2 = Document(DOCX_FILE)
    print(f"Paragraphs: {len(doc2.paragraphs)}")
    print(f"Tables: {len(doc2.tables)}")
    headings = [p for p in doc2.paragraphs if p.style.name.startswith("Heading")]
    print(f"Headings: {len(headings)}")


if __name__ == "__main__":
    main()
